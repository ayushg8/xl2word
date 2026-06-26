from docx import Document
from docx.shared import Inches
from xl2word.verify import detect_overflow

def test_detect_overflow_flags_wide_table(tmp_path):
    doc = Document()
    section = doc.sections[0]
    usable_in = (section.page_width - section.left_margin - section.right_margin) / 914400
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).width = Inches(usable_in)         # each col = full usable width
    t.cell(0, 1).width = Inches(usable_in)         # total = 2x usable -> overflow
    p = str(tmp_path / "wide.docx")
    doc.save(p)
    issues = detect_overflow(p)
    assert any("wider than" in s.lower() for s in issues)

def test_no_overflow_when_narrow(tmp_path):
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).width = Inches(1)
    t.cell(0, 1).width = Inches(1)
    p = str(tmp_path / "narrow.docx")
    doc.save(p)
    assert detect_overflow(p) == []

def test_no_false_overflow_on_merged_header(tmp_path):
    # Row 0 is a single banner merged across all 3 columns, but each grid column
    # fits the page. A spanned banner cell reports the full merged width, so reading
    # row cells would overcount ~3x; reading the grid must not false-positive.
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    for col in t.columns:
        col.width = Inches(1.5)            # grid: 3 x 1.5 = 4.5in within 6in usable
    for j in range(3):
        t.cell(1, j).width = Inches(1.5)   # data row fits its columns
    t.cell(0, 0).merge(t.cell(0, 2))       # banner spans every column
    p = str(tmp_path / "banner.docx")
    doc.save(p)
    assert detect_overflow(p) == []
