import os
from docx import Document
from xl2word.cli import main


def test_end_to_end_rich_sheet(build_rich_xlsx, tmp_path):
    src = build_rich_xlsx()                      # merged header + fill + percent + image
    out = str(tmp_path / "rich.docx")
    rc = main([src, "-o", out, "--workdir", str(tmp_path / "wd"), "--no-render"])
    assert rc == 0
    doc = Document(out)
    assert len(doc.tables) == 1
    flat = [c.text for row in doc.tables[0].rows for c in row.cells]
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "ESS LFP Cell Design" in heading_texts  # banner promoted to a section heading
    assert "1.2%" in flat                        # percent display survived end-to-end
    # the embedded image was sweep-extracted and placed
    assert len(doc.inline_shapes) >= 1
