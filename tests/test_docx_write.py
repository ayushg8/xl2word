import os
from docx import Document
from xl2word.extract import extract_semantic
from xl2word.layout import default_layout
from xl2word.docx_write import write_docx

def test_write_docx_builds_editable_table(build_simple_xlsx, tmp_path):
    wb = extract_semantic(build_simple_xlsx())
    out = str(tmp_path / "out.docx")
    write_docx(wb, default_layout(wb), out, images_dir=str(tmp_path / "images"))
    assert os.path.exists(out)
    doc = Document(out)
    assert len(doc.tables) == 1
    t = doc.tables[0]
    flat = [c.text for row in t.rows for c in row.cells]
    assert "Loading Level" in flat
    assert "25.40" in flat            # display value preserved
    assert "Parameter" in flat

def test_merged_header_is_merged(build_rich_xlsx, tmp_path):
    wb = extract_semantic(build_rich_xlsx(with_image=False))
    out = str(tmp_path / "rich.docx")
    write_docx(wb, default_layout(wb), out, images_dir=str(tmp_path / "images"))
    doc = Document(out)
    row0 = doc.tables[0].rows[0]
    # A1:C1 merged -> the three grid cells resolve to one underlying cell
    assert row0.cells[0]._tc is row0.cells[2]._tc

def test_normal_style_binds_eastasia_font(build_simple_xlsx, tmp_path):
    from docx.oxml.ns import qn
    wb = extract_semantic(build_simple_xlsx())
    out = str(tmp_path / "ea.docx")
    write_docx(wb, default_layout(wb), out, images_dir=str(tmp_path / "images"))
    style_el = Document(out).styles["Normal"].element
    rpr = style_el.find(qn("w:rPr"))
    assert rpr is not None
    rfonts = rpr.find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "Noto Sans CJK SC"

def test_orientation_change_starts_new_landscape_section(build_simple_xlsx, tmp_path):
    from docx.enum.section import WD_ORIENT
    from xl2word.layout import LayoutPlan, Block
    wb = extract_semantic(build_simple_xlsx())
    plan = LayoutPlan(title="t", blocks=[
        Block(kind="table", sheet="Spec", orientation="portrait"),
        Block(kind="table", sheet="Spec", orientation="landscape"),
    ])
    out = str(tmp_path / "rot.docx")
    write_docx(wb, plan, out, images_dir=str(tmp_path / "images"))
    doc = Document(out)
    assert len(doc.sections) >= 2
    assert doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE

def test_all_portrait_layout_stays_single_section(build_simple_xlsx, tmp_path):
    from docx.enum.section import WD_ORIENT
    from xl2word.layout import LayoutPlan, Block
    wb = extract_semantic(build_simple_xlsx())
    plan = LayoutPlan(title="t", blocks=[
        Block(kind="table", sheet="Spec", orientation="portrait"),
        Block(kind="table", sheet="Spec", orientation="portrait"),
    ])
    out = str(tmp_path / "port.docx")
    write_docx(wb, plan, out, images_dir=str(tmp_path / "images"))
    doc = Document(out)
    assert len(doc.sections) == 1
    assert doc.sections[0].orientation == WD_ORIENT.PORTRAIT
