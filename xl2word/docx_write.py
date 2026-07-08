from __future__ import annotations
import os
from docx import Document
from docx.shared import Emu, Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .model import Workbook, Sheet, Cell
from .layout import LayoutPlan, Block
from . import fit
from .cleaners import sanitize_authored

_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}
_DEFAULT_FONT = "Noto Sans CJK SC"   # renders Latin + Hangul; falls back if absent
_INK = "202124"           # near-black body text
_H1 = "1F3864"            # deep blue for sheet titles
_H2 = "2E5496"            # medium blue for section titles
_HEADER_FILL = "D9E1F2"   # light blue table header shade
_BORDER = "BFBFBF"        # light grey table borders


def _style_font(style, *, name=_DEFAULT_FONT, size=None, bold=None, color=None):
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), name)


_MARGIN = Inches(0.6)   # tight margins: this is dense documentation, not a report


def _tighten_margins(section) -> None:
    section.top_margin = section.bottom_margin = _MARGIN
    section.left_margin = section.right_margin = _MARGIN


def _new_document() -> Document:
    doc = Document()
    _tighten_margins(doc.sections[0])
    _style_font(doc.styles["Normal"], size=9, color=_INK)
    doc.styles["Normal"].paragraph_format.space_after = Pt(2)
    # Compact heading scale -- section labels, not report chapter titles.
    for name, size, color, sb, sa in (("Heading 1", 12, _H1, 8, 3),
                                      ("Heading 2", 9.5, _H2, 5, 2)):
        if name in doc.styles:
            _style_font(doc.styles[name], size=size, bold=True, color=color)
            pf = doc.styles[name].paragraph_format
            pf.space_before = Pt(sb)
            pf.space_after = Pt(sa)
            pf.keep_with_next = True
    return doc


def _section_for(doc, orientation: str):
    """Return a section in the requested orientation, starting a new section only
    when the current one's orientation differs. Keeps an all-portrait document to
    a single section while honoring per-block landscape requests."""
    section = doc.sections[-1]
    want_landscape = orientation == "landscape"
    if want_landscape == (section.orientation == WD_ORIENT.LANDSCAPE):
        return section
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    target = WD_ORIENT.LANDSCAPE if want_landscape else WD_ORIENT.PORTRAIT
    if section.orientation != target:
        section.orientation = target
        section.page_width, section.page_height = section.page_height, section.page_width
    _tighten_margins(section)
    return section


def _usable_width_emu(section) -> int:
    return int(section.page_width - section.left_margin - section.right_margin)


def _strip_extra_paragraphs(cell) -> None:
    """Remove trailing empty paragraphs left by a cell merge."""
    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    while len(paras) > 1:
        last = paras[-1]
        if not "".join(t.text or "" for t in last.findall(".//" + qn("w:t"))):
            tc.remove(last)
            paras = tc.findall(qn("w:p"))
        else:
            break


def _shade(cell, rgb: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), rgb)
    cell._tc.get_or_add_tcPr().append(shd)


# CT_TblPr children must appear in this schema order, or Word/LibreOffice ignore
# the out-of-order ones. In particular a misplaced tblLayout is dropped, and the
# table silently falls back to autofit, crushing short-content columns.
_TBLPR_ORDER = ("w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
                "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
                "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd", "w:tblLayout",
                "w:tblCellMar", "w:tblLook", "w:tblCaption", "w:tblDescription")


def _tblpr_set(table, el) -> None:
    """Insert an element into tblPr at its schema-correct position (replacing any
    existing element of the same tag)."""
    tblPr = table._tbl.tblPr
    tag = el.tag
    for existing in tblPr.findall(tag):
        tblPr.remove(existing)
    name = tag.split("}")[-1]
    after = _TBLPR_ORDER[_TBLPR_ORDER.index("w:" + name) + 1:]
    tblPr.insert_element_before(el, *after)


def _fixed_layout(table, total_twips: int) -> None:
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    _tblpr_set(table, layout)
    w = OxmlElement("w:tblW"); w.set(qn("w:type"), "dxa"); w.set(qn("w:w"), str(total_twips))
    _tblpr_set(table, w)


def _light_borders(table) -> None:
    """Thin light-grey grid instead of the heavy black Table Grid look."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER)
        borders.append(el)
    _tblpr_set(table, borders)


def _cell_padding(table, top=14, bottom=14, left=55, right=55) -> None:
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    _tblpr_set(table, mar)


def _style_header_row(row) -> None:
    """Bold the header row. Shade plain header cells light blue, but keep any
    source fill (e.g. coloured Anode/Cathode group headers) with legible contrast."""
    for cell in row.cells:
        fill = _cell_fill(cell)
        if not fill or fill.lower() == "auto":
            _shade(cell, _HEADER_FILL)
            fill = _HEADER_FILL
        ink = _ink_for(fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(ink)


def _bookmark(paragraph, name, bid) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bid))
    p = paragraph._p
    ppr = p.find(qn("w:pPr"))
    p.insert(list(p).index(ppr) + 1 if ppr is not None else 0, start)
    p.append(end)


def _hyperlink(paragraph, anchor, text, color) -> None:
    link = OxmlElement("w:hyperlink"); link.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), color); rpr.append(col)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t); link.append(r)
    paragraph._p.append(link)


def _hyperlink_run(paragraph, anchor, text, color, size_pt) -> None:
    """Append an inline hyperlink run to an existing paragraph (for the compact
    one-line contents strip)."""
    link = OxmlElement("w:hyperlink"); link.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), color); rpr.append(col)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t); link.append(r)
    paragraph._p.append(link)


def _page_number_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, text in (("begin", None), ("instr", " PAGE "), ("end", None)):
        if kind == "instr":
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = text
            run._r.append(it)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)


def _set_cell_width(cell, width_emu: int) -> None:
    cell.width = Emu(width_emu)


def _grid(sheet: Sheet, region):
    if region:
        r0, c0, r1, c1 = region
    else:
        r0, c0, r1, c1 = 1, 1, sheet.max_row, sheet.max_col
    by_pos = {(c.row, c.col): c for c in sheet.cells}
    rows = []
    for r in range(r0, r1 + 1):
        rows.append([by_pos.get((r, c)) for c in range(c0, c1 + 1)])
    return rows, (r0, c0, r1, c1)


def _ink_for(fill_hex: str | None) -> str:
    """Legible text colour for a given cell fill: white on a dark fill, near-black
    otherwise. Replaces blanket dark text, which turned white-on-navy banners into
    illegible dark-on-dark."""
    if not fill_hex or fill_hex.lower() in ("auto", "none"):
        return _INK
    try:
        r, g, b = int(fill_hex[0:2], 16), int(fill_hex[2:4], 16), int(fill_hex[4:6], 16)
    except (ValueError, IndexError):
        return _INK
    return "FFFFFF" if (0.299 * r + 0.587 * g + 0.114 * b) < 140 else _INK


def _cell_fill(cell) -> str | None:
    tcpr = cell._tc.find(qn("w:tcPr"))
    shd = tcpr.find(qn("w:shd")) if tcpr is not None else None
    return shd.get(qn("w:fill")) if shd is not None else None


def _cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit"); el.set(qn("w:val"), "true")
    trPr.append(el)


def _keep_next(row) -> None:
    """Keep this row with the following one, so the whole table stays on one page."""
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = True


def _apply_cell(docx_cell, model_cell: Cell | None, font_pt: float = 10) -> None:
    para = docx_cell.paragraphs[0]
    # Normalise em/en dashes to hyphens even in source cells: a typographic rule for
    # the whole document (no em dashes), and it never touches numbers or values.
    text = (model_cell.display or "").replace("—", " - ").replace("–", "-") if model_cell else ""
    run = para.add_run(text)
    run.font.size = Pt(font_pt)
    if model_cell:
        st = model_cell.style
        run.bold = st.bold
        run.italic = st.italic
        run.font.color.rgb = RGBColor.from_string(_ink_for(st.fill))  # contrast, not blanket dark
        if st.align_h in _ALIGN:
            para.alignment = _ALIGN[st.align_h]
        if st.fill:
            _shade(docx_cell, st.fill)
    else:
        run.font.color.rgb = RGBColor.from_string(_INK)


def _pruned_axes(sheet: Sheet, rows, r0, c0, r1, c1):
    """Return the original row and column numbers to keep after dropping rows and
    columns that are fully empty within the region. A row or column that any merge
    covers is always kept, so merges stay contiguous and remap cleanly."""
    merges = [m for m in sheet.merged
              if m.min_row >= r0 and m.max_row <= r1 and m.min_col >= c0 and m.max_col <= c1]
    merge_rows, merge_cols = set(), set()
    for m in merges:
        merge_rows.update(range(m.min_row, m.max_row + 1))
        merge_cols.update(range(m.min_col, m.max_col + 1))

    def nonempty(cell):
        return bool(cell and (cell.display or "").strip())

    keep_rows = [r0 + gi for gi, row in enumerate(rows)
                 if (r0 + gi) in merge_rows or any(nonempty(c) for c in row)]
    keep_cols = [c0 + gj for gj in range(c1 - c0 + 1)
                 if (c0 + gj) in merge_cols or any(nonempty(rows[gi][gj]) for gi in range(len(rows)))]
    return keep_rows or [r0], keep_cols or [c0], merges


_EMU_IN = 914400
_PORTRAIT_USABLE = _EMU_IN * 73 // 10    # 8.5in page - 2*0.6 margin = 7.3in
_LANDSCAPE_USABLE = _EMU_IN * 98 // 10   # 11in page - 2*0.6 margin = 9.8in
_PORTRAIT_HEIGHT = _EMU_IN * 95 // 10    # 11in - 2*0.6 margin - footer ~= 9.5in
_LANDSCAPE_HEIGHT = _EMU_IN * 71 // 10   # 8.5in - 2*0.6 margin - footer ~= 7.1in
_MIN_FONT = 6                            # never shrink table text below this


def _header_band(keep_rows, merges, nrows) -> int:
    """How many leading rows form the header: rows spanned by a multi-column merge
    (a merged header hierarchy) plus the leaf-label row beneath them, else 1."""
    hmerge_rows = set()
    for m in merges:
        if m.max_col > m.min_col:
            hmerge_rows.update(range(m.min_row, m.max_row + 1))
    hdr = 0
    while hdr < nrows and keep_rows[hdr] in hmerge_rows:
        hdr += 1
    if 0 < hdr < nrows:
        hdr += 1
    return max(1, min(hdr, 4, nrows - 1 if nrows > 1 else 1))


def _looks_numeric(text: str) -> bool:
    core = text.strip().lstrip("<>=≤≥±~").rstrip("%").replace(",", "").replace("±", " ").split()
    try:
        float(core[0]) if core else float("nan")
        return bool(core)
    except ValueError:
        return False


def _label_col_count(rows) -> int:
    """Leading columns that identify rows (mostly text, e.g. a parameter name),
    which are repeated in every chunk when a wide table is split by columns."""
    n = 0
    for gj in range(min(2, len(rows[0]) if rows else 0)):
        vals = [(row[gj].display if gj < len(row) and row[gj] else "").strip() for row in rows]
        nonempty = [v for v in vals if v]
        if not nonempty:
            break
        numeric = sum(1 for v in nonempty if _looks_numeric(v))
        if numeric <= len(nonempty) * 0.4:
            n += 1
        else:
            break
    return max(1, n)


_BASE_FONT = 8   # compact base size for table text


def _fit_font(text_rows, usable_w, usable_h) -> tuple:
    """Pick the largest font (<= base, >= min) at which the table fits the page in
    both width and height, and return (font_pt, column widths)."""
    font_pt = _BASE_FONT
    min_sum = fit.min_width_sum(text_rows, font_pt)
    if min_sum > usable_w:                       # too wide even at base: shrink for width
        font_pt = max(_MIN_FONT, _BASE_FONT * usable_w // min_sum)
    widths = fit.balanced_column_widths(text_rows, usable_w, font_pt)
    for _ in range(6):                           # shrink for height until it fits one page
        h = fit.estimate_table_height_emu(text_rows, widths, font_pt)
        if h <= usable_h or font_pt <= _MIN_FONT:
            break
        font_pt = max(_MIN_FONT, int(font_pt * usable_h / h))
        widths = fit.balanced_column_widths(text_rows, usable_w, font_pt)
    return font_pt, widths


def _render_table(doc, full_rows, r0, c0, keep_rows, keep_cols, merges,
                  orientation, hdr_count, caption=None) -> None:
    section = _section_for(doc, orientation)
    row_map = {orig: i for i, orig in enumerate(keep_rows)}
    col_map = {orig: j for j, orig in enumerate(keep_cols)}
    rows = [[full_rows[orig - r0][cc - c0] for cc in keep_cols] for orig in keep_rows]
    nrows, ncols = len(rows), len(keep_cols)
    text_rows = [[(cell.display if cell else "") for cell in row] for row in rows]
    usable = _usable_width_emu(section)
    usable_h = _LANDSCAPE_HEIGHT if orientation == "landscape" else _PORTRAIT_HEIGHT
    font_pt, widths = _fit_font(text_rows, usable, usable_h)

    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(1); cp.paragraph_format.space_after = Pt(1)
        run = cp.add_run(sanitize_authored(caption)); run.italic = True; run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string(_H2)

    fits_one_page = fit.estimate_table_height_emu(text_rows, widths, font_pt) <= usable_h

    table = doc.add_table(rows=nrows, cols=ncols)
    _light_borders(table)
    _cell_padding(table)
    _fixed_layout(table, sum(int(w / 635) for w in widths))
    for gi, row in enumerate(rows):
        for gj, mcell in enumerate(row):
            _apply_cell(table.cell(gi, gj), mcell, font_pt)
            _set_cell_width(table.cell(gi, gj), widths[gj])
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w_emu in zip(grid.findall(qn("w:gridCol")), widths):
            gc.set(qn("w:w"), str(int(w_emu / 635)))
    # Rows never split mid-row. If the whole table fits one page, also keep every
    # row with the next so Word holds the table together on a single page (no
    # cross-page continuation). If it is genuinely taller than a page even at the
    # min font, let it flow so it packs cleanly instead of stranding a few rows.
    for i, row in enumerate(table.rows):
        _cant_split(row)
        if fits_one_page and i < nrows - 1:
            _keep_next(row)
    # Apply only merges whose whole span is inside this column set (a chunk may
    # exclude some columns a merge covers; skip those to keep the grid valid).
    for m in merges:
        if (m.min_row in row_map and m.max_row in row_map
                and all(cc in col_map for cc in range(m.min_col, m.max_col + 1))):
            a = table.cell(row_map[m.min_row], col_map[m.min_col])
            b = table.cell(row_map[m.max_row], col_map[m.max_col])
            a.merge(b)
            _strip_extra_paragraphs(a)
    for gi in range(min(hdr_count, nrows)):
        if nrows >= 2:
            _style_header_row(table.rows[gi])


def _add_table(doc, sheet: Sheet, block: Block) -> None:
    full_rows, (r0, c0, r1, c1) = _grid(sheet, block.region)
    if len(full_rows) == 0 or (c1 - c0 + 1) == 0:
        return
    keep_rows, keep_cols, merges = _pruned_axes(sheet, full_rows, r0, c0, r1, c1)
    rows = [[full_rows[orig - r0][cc - c0] for cc in keep_cols] for orig in keep_rows]
    text_rows = [[(cell.display if cell else "") for cell in row] for row in rows]
    hdr = _header_band(keep_rows, merges, len(keep_rows))
    usable = _LANDSCAPE_USABLE if block.orientation == "landscape" else _PORTRAIT_USABLE

    # If the table fits horizontally (even at the min font), render it as one table.
    if fit.min_width_sum(text_rows, _MIN_FONT) <= usable:
        _render_table(doc, full_rows, r0, c0, keep_rows, keep_cols, merges, block.orientation, hdr)
        return

    # Too wide for the page: split the data columns into chunks that each fit,
    # repeating the label column(s) and header rows, so no column is lost off-page.
    label_n = _label_col_count(rows)
    label_cols = keep_cols[:label_n]
    data_cols = keep_cols[label_n:]

    def width_of(cols):
        sub = [[(full_rows[orig - r0][cc - c0].display if full_rows[orig - r0][cc - c0] else "")
                for cc in cols] for orig in keep_rows]
        return fit.min_width_sum(sub, _MIN_FONT)

    chunks, cur = [], []
    for dc in data_cols:
        if cur and width_of(label_cols + cur + [dc]) > usable:
            chunks.append(cur); cur = [dc]
        else:
            cur.append(dc)
    if cur:
        chunks.append(cur)

    pos, ncols = 0, len(keep_cols)
    for ci, chunk in enumerate(chunks):
        cols = label_cols + chunk
        cap = (f"Columns {label_n + pos + 1}–{label_n + pos + len(chunk)} of {ncols}"
               if len(chunks) > 1 else None)
        _render_table(doc, full_rows, r0, c0, keep_rows, cols, merges, block.orientation, hdr, cap)
        pos += len(chunk)
        if ci < len(chunks) - 1:
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)


def _repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _add_prose(doc, text: str) -> None:
    """A short plain-English intro line under a heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(sanitize_authored(text))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(_INK)


def _add_bullets(doc, items: list) -> None:
    """A real bulleted list -- for notes/instructions that are a list, not a grid."""
    for it in items:
        text = sanitize_authored(it if isinstance(it, str) else str(it))
        if not text:
            continue
        p = doc.add_paragraph(style="List Bullet")
        pf = p.paragraph_format
        pf.space_after = Pt(1); pf.space_before = Pt(0)
        pf.left_indent = Inches(0.25); pf.first_line_indent = Inches(-0.13)
        run = p.add_run(text)
        run.font.size = Pt(9); run.font.color.rgb = RGBColor.from_string(_INK)


def _add_note(doc, text: str) -> None:
    """A subtle callout for a single important note (indented, shaded, italic)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.1); pf.space_before = Pt(2); pf.space_after = Pt(3)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F1F4FA")
    p._p.get_or_add_pPr().append(shd)
    run = p.add_run(sanitize_authored(text))
    run.italic = True; run.font.size = Pt(8.5); run.font.color.rgb = RGBColor.from_string(_H2)


def _add_keyvalue(doc, pairs: list) -> None:
    """A compact settings block: bold key, aligned value -- for a small set of
    parameter/value settings that read better as a definition list than a grid."""
    valid = [(str(k).strip(), str(v).strip()) for k, v in pairs if str(k).strip()]
    if not valid:
        return
    key_w = min(_EMU_IN * 32 // 10, max((fit._text_width_emu(len(k), 9) for k, _ in valid), default=0))
    table = doc.add_table(rows=len(valid), cols=2)
    _cell_padding(table, top=6, bottom=6, left=0, right=60)
    _fixed_layout(table, int((key_w + _EMU_IN * 24 // 10) / 635))
    for i, (k, v) in enumerate(valid):
        kc, vc = table.cell(i, 0), table.cell(i, 1)
        _set_cell_width(kc, key_w); _set_cell_width(vc, _EMU_IN * 24 // 10)
        kr = kc.paragraphs[0].add_run(sanitize_authored(k))
        kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = RGBColor.from_string(_INK)
        vr = vc.paragraphs[0].add_run(v)
        vr.font.size = Pt(9); vr.font.color.rgb = RGBColor.from_string(_INK)
        for c in (kc, vc):
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for i, row in enumerate(table.rows):
        _cant_split(row)
        if i < len(table.rows) - 1:
            _keep_next(row)


def _add_image(doc, images_dir: str, block: Block) -> None:
    from docx.shared import Inches
    path = block.path
    candidate = path if os.path.isabs(path) else os.path.join(os.path.dirname(images_dir), path)
    if not os.path.exists(candidate):
        candidate = os.path.join(images_dir, os.path.basename(path))
    if os.path.exists(candidate):
        doc.add_picture(candidate, width=Inches(5))
        if block.caption:
            cap = doc.add_paragraph(block.caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_docx(wb: Workbook, layout: LayoutPlan, out_path: str, images_dir: str) -> None:
    doc = _new_document()
    by_name = {s.name: s for s in wb.sheets}
    # Assign a bookmark anchor to every heading so the contents can link to it.
    anchors = {i: f"sec{i}" for i, b in enumerate(layout.blocks) if b.kind == "heading"}

    # Compact title block (no cover page): title line, then a tight one-line
    # contents of the sheets, then straight into content. This is documentation,
    # not a report -- no ceremony.
    if layout.title:
        t = doc.add_paragraph()
        t.paragraph_format.space_after = Pt(2)
        run = t.add_run(layout.title)
        run.bold = True; run.font.size = Pt(15); run.font.color.rgb = RGBColor.from_string(_H1)
    sheet_headings = [(i, b) for i, b in enumerate(layout.blocks)
                      if b.kind == "heading" and b.level == 1 and (b.text or "").strip()]
    if len(sheet_headings) > 1:
        toc = doc.add_paragraph()
        toc.paragraph_format.space_after = Pt(8)
        lbl = toc.add_run("Contents:  ")
        lbl.bold = True; lbl.font.size = Pt(8); lbl.font.color.rgb = RGBColor.from_string(_H2)
        for n, (i, b) in enumerate(sheet_headings):
            if n:
                sep = toc.add_run("   ·   "); sep.font.size = Pt(8)
                sep.font.color.rgb = RGBColor.from_string(_BORDER)
            _hyperlink_run(toc, anchors[i], b.text, _H2, 8)

    prev_kind = None
    for idx, block in enumerate(layout.blocks):
        if block.kind == "heading":
            # Put the heading in the same section orientation as the table it
            # introduces, scanning past any intervening headings, so a run of
            # headings above a landscape table is not stranded on a portrait page.
            for later in layout.blocks[idx + 1:]:
                if later.kind == "table":
                    _section_for(doc, later.orientation)
                    break
                if later.kind == "pagebreak":
                    break
            h = doc.add_heading(sanitize_authored(block.text or ""), level=block.level or 1)
            _bookmark(h, anchors[idx], idx)
        elif block.kind == "table" and block.sheet in by_name:
            # Word merges two tables that are not separated by a paragraph, which
            # scrambles the combined column grid (columns collapse to slivers). Emit
            # a thin spacer paragraph between back-to-back tables to keep them apart.
            if prev_kind == "table":
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after = Pt(4)
            _add_table(doc, by_name[block.sheet], block)
        elif block.kind == "prose" and block.text:
            _add_prose(doc, block.text)
        elif block.kind == "bullets" and block.items:
            _add_bullets(doc, block.items)
        elif block.kind == "note" and block.text:
            _add_note(doc, block.text)
        elif block.kind == "keyvalue" and block.pairs:
            _add_keyvalue(doc, block.pairs)
        elif block.kind == "image" and block.path:
            _add_image(doc, images_dir, block)
        elif block.kind == "pagebreak":
            doc.add_page_break()
        prev_kind = block.kind
    for section in doc.sections:
        _page_number_footer(section)
    doc.save(out_path)
